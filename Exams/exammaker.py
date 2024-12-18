import csv
import random
from docx import Document
from docx.shared import Pt
import math

def is_within_sections(question_id, min_section, max_section):
    # Extract the section number from the question ID and check if it is within the given range
    section = int(question_id.split('.')[0])
    return min_section <= section <= max_section

def main():
    # Read the questions from the CSV file
    with open('questions_answers_all.csv', 'r', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        questions_list = [question for question in reader if is_within_sections(question['ID'], 5, 9)]


    total_questions = len(questions_list)
    if total_questions == 0:
        print("No questions found in the CSV file.")
        return

    # Shuffle the list to randomize question order
    random.shuffle(questions_list)

    # Determine the number of documents needed
    num_docs = math.ceil(total_questions / 50)

    for doc_num in range(1, num_docs + 1):
        # Calculate the start and end indices for the current batch
        start_idx = (doc_num - 1) * 50
        end_idx = min(start_idx + 50, total_questions)

        # Get the batch of questions for this document
        batch_questions = questions_list[start_idx:end_idx]

        # Create a new Word Document
        document = Document()

        # Set default font size and style
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        # Write the questions to the document
        document.add_heading(f'Randomly Selected Questions - Set {doc_num}', level=1)
        for idx, question in enumerate(batch_questions, 1):
            p = document.add_paragraph()
            p.add_run(f"Question {idx} (ID: {question['ID']}): ").bold = True
            p.add_run(question['Question'])

        # Add a page break
        document.add_page_break()

        # Write the answers at the end
        document.add_heading('Answers', level=1)
        for idx, question in enumerate(batch_questions, 1):
            p = document.add_paragraph()
            p.add_run(f"Answer to Question {idx}: ").bold = True
            p.add_run(question['Answer'])

        # Save the document
        doc_filename = f'practice/exam2/Random_Questions_Set_{doc_num}.docx'
        document.save(doc_filename)
        print(f"Document '{doc_filename}' has been created successfully.")

if __name__ == "__main__":
    main()
